from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "Micene_Step12_Prototipo_tecnico_controllato_PWA_scena_P0_PLACEHOLDER.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667784"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
CYAN = "1E8094"
MAGENTA = "A51A64"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, color: str, size=8, space=6, side="bottom") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def font(run, name="Calibri", size=None, bold=None, italic=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_keep(paragraph, next_=False, together=False) -> None:
    paragraph.paragraph_format.keep_with_next = next_
    paragraph.paragraph_format.keep_together = together


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def make_table(doc, headers, rows, widths):
    assert sum(widths) == 9360
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    def format_row(row, values, header=False):
        for index, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if header:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            font(run, size=9.2, bold=header, color=INK if header else None)
            if index == 0 and not header:
                run.bold = True

    format_row(table.rows[0], headers, header=True)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for values in rows:
        format_row(table.add_row(), values)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_caption(doc, text: str):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    set_keep(p, next_=False, together=True)
    return p


def add_picture_with_alt(doc, path: Path, width_inches: float, alt: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width_inches))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)
    c_nv_pr = shape._inline.graphic.graphicData.pic.nvPicPr.cNvPr
    c_nv_pr.set("descr", alt)
    set_keep(p, next_=True, together=True)
    return p


def add_callout(doc, label: str, text: str, color=CYAN):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.08
    set_paragraph_shading(p, CALLOUT)
    set_paragraph_border(p, color, size=18, space=6, side="left")
    r = p.add_run(f"{label}  ")
    font(r, bold=True, color=color)
    font(p.add_run(text), color=INK)
    set_keep(p, together=True)


def add_body(doc, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), bold=True, color=INK)
        font(p.add_run(text[len(bold_lead):]))
    else:
        p.add_run(text)
    return p


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)

    if "Kicker" not in doc.styles:
        kicker = doc.styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        kicker = doc.styles["Kicker"]
    kicker.font.name = "Calibri"
    kicker.font.size = Pt(10)
    kicker.font.bold = True
    kicker.font.color.rgb = RGBColor.from_string(CYAN)
    kicker.paragraph_format.space_after = Pt(8)
    kicker.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    font(header.add_run("MICENE — STEP 12"), size=8.5, bold=True, color=MUTED)
    font(header.add_run("\tPROTOTIPO TECNICO CONTROLLATO"), size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    font(footer.add_run("P0-MIN CHIUSO · PLACEHOLDER"), size=8.5, color=MUTED)
    font(footer.add_run("\tPagina "), size=8.5, color=MUTED)
    add_page_field(footer)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    package = json.loads((ROOT / "package.json").read_text())
    deps = {**package.get("dependencies", {}), **package.get("devDependencies", {})}
    runtime = json.loads((ROOT / "reports/performance/runtime-metrics.json").read_text())
    build = json.loads((ROOT / "reports/performance/build-sizes.json").read_text())
    metric_map = {item["label"]: item["value"] for item in runtime["metrics"]}
    babylon = next(item for item in build["files"] if item["file"].startswith("assets/babylon-") and item["file"].endswith(".js"))

    doc = Document()
    style_document(doc)
    doc.core_properties.title = "Micene — Step 12 — Prototipo tecnico controllato della PWA e scena P0 PLACEHOLDER"
    doc.core_properties.subject = "Relazione tecnica conclusiva dello Step 12"
    doc.core_properties.author = "Progetto Micene"
    doc.core_properties.keywords = "Micene, PWA, Babylon.js, PLACEHOLDER, Step 12"

    p = doc.add_paragraph(style="Kicker")
    p.add_run("MICENE · ESPLORAZIONE ARCHEOLOGICA 3D")
    p = doc.add_paragraph(style="Title")
    p.add_run("Micene — Step 12 — Prototipo tecnico controllato della PWA e scena P0 PLACEHOLDER")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Relazione di implementazione, verifica e continuità · 21 luglio 2026")
    for label, value in (
        ("Stato", "Implementazione completata e verificata nell’ambiente disponibile"),
        ("Gate tecnici", "T1 PASS · T2 PASS · base T4 PASS STRUTTURALE"),
        ("Gate archeologico", "P0-MIN CHIUSO"),
        ("Dispositivo reale", "Safari/iPad PENDING DEVICE"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        font(p.add_run(f"{label}: "), bold=True, color=INK)
        font(p.add_run(value))
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_border(rule, CYAN, size=16, space=3)

    add_callout(doc, "ESITO", "È stato costruito un software realmente eseguibile. La scena contiene esclusivamente primitive tecniche astratte; nessuna geometria è presentata come misura o ricostruzione di Micene.")

    doc.add_heading("1. Esito conseguito", level=1)
    add_body(doc, "La PWA dispone di home, esploratore 3D, fallback 2D, manifest installabile, service worker, aggiornamento controllato, stato offline, messaggi d’errore e retry. La lingua iniziale è l’italiano; l’inglese è predisposto nella struttura i18n.")
    add_body(doc, "Babylon.js è incapsulato dietro un EngineAdapter. Lo store Zustand conserva soltanto posa, modalità, selezione, lingua e profilo grafico come record serializzabili; nessun oggetto del motore entra nei dati scientifici.")
    make_table(doc,
        ["Gate", "Esito", "Evidenza", "Limite"],
        [
            ["T1", "PASS", "Build, PWA, offline, fallback, axe", "iPad reale PENDING DEVICE"],
            ["T2", "PASS", "Scena, camere, input, picking, diagnostica", "iPad reale PENDING DEVICE"],
            ["Base T4", "PASS STRUTTURALE", "featureId → scheda; legenda e NULL espliciti", "Non è T4 scientifico completo"],
            ["P0-MIN", "CHIUSO", "Invariato", "Mancano dati metrici e documentari"],
        ], [1000, 1600, 3520, 3240])

    doc.add_page_break()
    doc.add_heading("2. Implementazione reale", level=1)
    make_table(doc,
        ["Componente", "Realizzato", "Controllo epistemico"],
        [
            ["Shell PWA", "Route, responsive UI, manifest, SW, update prompt", "Nessuna telemetria o CDN runtime"],
            ["Scena", "Canvas WebGL2, griglia, assi, piano, luce e volumi", "Solo DEBUG_DISPLAY; M0"],
            ["Navigazione", "Prima/terza persona, WASD, touch, reset, limiti", "Collider e limiti soltanto tecnici"],
            ["Selezione", "Picking Babylon e persistenza per featureId", "Nomi mesh non usati come identità"],
            ["Modalità scientifica", "Scheda HTML, legenda A–D/M0–M3 e natura epistemica", "PLACEHOLDER fuori da A–D"],
            ["Validazione", "Zod, JSON Schema e guardia di build", "Blocca definitivo/released e asset vietati"],
            ["Fallback", "Pagina 2D accessibile e retry del manifest", "Nessuna falsa ricostruzione statica"],
        ], [1800, 3540, 4020])

    doc.add_heading("3. Struttura del repository", level=1)
    code = doc.add_paragraph()
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(10)
    code.paragraph_format.left_indent = Inches(0.18)
    set_paragraph_shading(code, "EEF3F6")
    tree = (
        "micene-pwa/\n"
        "  apps/web/src/{app,components,engine,features,services,state,styles}/\n"
        "  assets/{placeholders,release,manifests}/\n"
        "  data/{assets,schemas,scientific,sources,decisions}/\n"
        "  tests/e2e/ · tools/ · reports/ · screenshots/\n"
        "  docs/{architecture,scientific-method,operations,gates}/\n"
        "  registries/ · dist/\n"
        "  PROJECT_STATE.md · DECISIONS.md · NEXT_STEP.md"
    )
    font(code.add_run(tree), name="Consolas", size=9, color=INK)

    doc.add_page_break()
    doc.add_heading("4. Versioni effettivamente installate", level=1)
    version_rows = [
        ["React / React DOM", deps["react"].lstrip("^~"), "Shell e componenti"],
        ["React Router DOM", deps["react-router-dom"].lstrip("^~"), "Routing client"],
        ["Vite", deps["vite"].lstrip("^~"), "Build"],
        ["TypeScript", deps["typescript"].lstrip("^~"), "Strict; ultima versione compatibile con lint"],
        ["Babylon.js Core", deps["@babylonjs/core"].lstrip("^~"), "Motore WebGL2"],
        ["Zustand", deps["zustand"].lstrip("^~"), "Stato serializzabile"],
        ["Zod", deps["zod"].lstrip("^~"), "Validazione runtime"],
        ["i18next / react-i18next", f"{deps['i18next'].lstrip('^~')} / {deps['react-i18next'].lstrip('^~')}", "Internazionalizzazione"],
        ["vite-plugin-pwa", deps["vite-plugin-pwa"].lstrip("^~"), "Workbox/PWA"],
        ["Vitest", deps["vitest"].lstrip("^~"), "Test unitari"],
        ["Playwright", deps["@playwright/test"].lstrip("^~"), "E2E"],
        ["axe-core", deps["axe-core"].lstrip("^~"), "Accessibilità"],
    ]
    make_table(doc, ["Pacchetto", "Versione", "Ruolo"], version_rows, [3100, 1500, 4760])
    add_callout(doc, "NOTA VERSIONI", "TypeScript 7.0.2 è stato provato ma risultava incompatibile con typescript-eslint 8.65.0. È stata quindi fissata la versione stabile 6.0.3; la decisione è registrata e riproducibile nel lockfile.", color=GOLD)

    doc.add_page_break()
    doc.add_heading("5. Test eseguiti e risultati", level=1)
    make_table(doc,
        ["Controllo", "Esito", "Risultato verificato"],
        [
            ["Registry/allowlist", "PASS", "10 record; nessun asset definitivo, personale o istituzionale"],
            ["ESLint", "PASS", "Nessun errore"],
            ["TypeScript strict", "PASS", "Nessun errore"],
            ["Vitest/Testing Library", "PASS", "4 file; 7/7 test"],
            ["Build produzione", "PASS", "Vite + manifest + service worker"],
            ["axe-core", "PASS", "0 violazioni su home, fallback ed explorer"],
            ["E2E desktop", "PASS", "Picking, camera, touch HTML, retry, offline"],
            ["Tablet verticale", "SIMULATED PASS", "834×1194; touch simulato"],
            ["Tablet orizzontale", "SIMULATED PASS", "1194×834; touch simulato"],
            ["Safari/iPad reale", "PENDING DEVICE", "Non eseguito"],
        ], [2600, 1700, 5060])

    doc.add_heading("6. Misure realmente ottenute", level=1)
    add_callout(doc, "AMBIENTE", f"Chromium {runtime['environment']['browser']} headless, WebGL2 tramite ANGLE/SwiftShader, viewport 1280×720. Questi valori non sono misure di un iPad reale.", color=GOLD)
    make_table(doc,
        ["Metrica", "Risultato", "Qualifica"],
        [
            ["FPS", metric_map["FPS"], "MISURATO in headless"],
            ["Frame time", metric_map["Frame time"], "MISURATO"],
            ["Draw call", metric_map["Draw call"], "MISURATO"],
            ["Triangoli visibili", metric_map["Triangoli visibili"], "MISURATO"],
            ["Backend", metric_map["Backend"], "RILEVATO"],
            ["Shell iniziale gzip", f"{build['initialShellGzipBytes']:,} byte".replace(",", "."), "CALCOLATO sulla build"],
            ["Precache Workbox", "2.080,21 KiB", "Sopra target 2 MiB; sotto soglia 4 MiB"],
            ["Chunk Babylon", f"{babylon['bytes']:,} byte".replace(",", "."), "Lazy-loaded sull’esploratore"],
        ], [2800, 1900, 4660])

    doc.add_page_break()
    doc.add_heading("7. Evidenze visive", level=1)
    add_picture_with_alt(doc, ROOT / "screenshots/01-home-desktop.png", 6.5, "Home del prototipo tecnico Micene con Gate e avvertenze PLACEHOLDER.")
    add_caption(doc, "Figura 1 · Home reale della build di produzione, viewport desktop.")

    doc.add_page_break()
    add_picture_with_alt(doc, ROOT / "screenshots/04-explorer-selection-desktop.png", 6.5, "Esploratore 3D con primitive wireframe e scheda PLACEHOLDER selezionata.")
    add_caption(doc, "Figura 2 · Picking per featureId e scheda scientifica strutturale.")
    add_body(doc, "La resa impiega colori diagnostici, wireframe e watermark persistente. Non sono presenti pietra, luce cinematografica, alzati realistici o simulazioni del 1250 a.C.")

    doc.add_page_break()
    add_picture_with_alt(doc, ROOT / "screenshots/05-explorer-tablet-portrait-simulated.png", 4.35, "Esploratore in viewport tablet verticale simulato.")
    add_caption(doc, "Figura 3 · Viewport 834×1194 con touch simulato; non è una prova su iPad reale.")

    doc.add_heading("8. Problemi residui", level=1)
    make_table(doc,
        ["ID", "Priorità", "Problema e conseguenza"],
        [
            ["I12-01", "Alta", "Installazione, aggiornamento, offline, touch e performance da verificare su Safari/iPad reale."],
            ["I12-02", "Bassa", "Precache circa 2,03 MiB: supera marginalmente il target, non la soglia."],
            ["I12-03", "Media", "Chunk Babylon ampio; ottimizzazione ulteriore dopo profiling reale."],
            ["I12-04", "Bloccante archeologico", "P0-MIN chiuso; nessun PLACEHOLDER può essere sostituito con geometria archeologica."],
        ], [1100, 2100, 6160])

    doc.add_heading("9. Stato finale e continuità", level=1)
    add_body(doc, "R01, R02, R03 e R05 restano pronte dopo compilazione o conferma e non inviate. R04 resta sospesa e non inviabile. Nessuna risposta, licenza o autorizzazione è considerata ricevuta; nessun dato personale è stato utilizzato.")
    add_callout(doc, "P0-MIN", "Resta CHIUSO. T1, T2 e la base strutturale T4 dimostrano il software, non la sufficienza archeologica o metrica.", color=RED)
    doc.add_heading("Unica attività successiva autorizzata", level=2)
    add_body(doc, "Step 13 — Collaudo controllato su iPad reale: registrare modello, iPadOS/Safari, installazione PWA, aggiornamento, offline, gesti touch, prima/terza persona, picking, FPS/frame time, errori e schermate. L’attività non autorizza import di dati archeologici, ricostruzione 1250, invio di richieste o apertura di P0-MIN.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
